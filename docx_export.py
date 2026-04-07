"""Экспорт кроссворда в DOCX (Word) через python-docx."""

import os
from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from models import Crossword


def export_docx(crossword: Crossword, filepath: str, show_answers: bool = False) -> None:
    """Экспортирует кроссворд в DOCX-файл."""
    doc = Document()

    # Настройка полей страницы
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    _add_crossword_page(doc, crossword, show_answers)
    doc.add_page_break()
    _add_clues_page(doc, crossword)

    doc.save(filepath)


def _add_crossword_page(doc: Document, cw: Crossword, show_answers: bool) -> None:
    """Добавляет страницу с сеткой кроссворда."""
    # Заголовок
    title = cw.title
    if show_answers:
        title += " (ОТВЕТЫ)"

    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    p = doc.add_paragraph("Кроссворд для воскресной школы")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # отступ

    # Собираем номера по позициям
    numbers: dict[tuple[int, int], int] = {}
    for w in cw.words:
        key = (w.row, w.col)
        if key not in numbers:
            numbers[key] = w.number

    # Вычисляем размер ячейки (чтобы поместиться на страницу)
    available_width = Mm(170)  # ~A4 минус поля
    cell_size = min(Mm(8), available_width / cw.cols)
    cell_size = max(cell_size, Mm(4))  # минимум

    # Таблица для сетки
    table = doc.add_table(rows=cw.rows, cols=cw.cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Убираем рамки по умолчанию
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tbl_pr.append(borders)

    for row_idx in range(cw.rows):
        row = table.rows[row_idx]
        row.height = cell_size

        for col_idx in range(cw.cols):
            cell = row.cells[col_idx]
            cell.width = cell_size

            ch = cw.grid[row_idx][col_idx]
            if ch:
                # Белая клетка с рамкой
                _set_cell_border(cell, "333333")
                _set_cell_shading(cell, "FFFFFF")

                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Номер
                num_key = (row_idx, col_idx)
                if num_key in numbers:
                    run_num = p.add_run(str(numbers[num_key]))
                    run_num.font.size = Pt(5)
                    run_num.font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)
                    run_num.font.superscript = True

                # Буква (если показываем ответы)
                if show_answers:
                    run_letter = p.add_run(f"\n{ch}")
                    run_letter.font.size = Pt(9)
                    run_letter.font.bold = True
                    run_letter.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

                # Уменьшаем отступы в ячейке
                _set_cell_margins(cell, top=20, bottom=20, left=40, right=40)
            else:
                # Серая/тёмная ячейка (пустая)
                _set_cell_shading(cell, "E0E0E0")
                _set_cell_border(cell, "E0E0E0")


def _add_clues_page(doc: Document, cw: Crossword) -> None:
    """Добавляет страницу с вопросами."""
    h = doc.add_heading("Вопросы к кроссворду", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    across = sorted([w for w in cw.words if w.direction == 'across'], key=lambda w: w.number)
    down = sorted([w for w in cw.words if w.direction == 'down'], key=lambda w: w.number)

    if across:
        _add_clue_section(doc, "По горизонтали →", across)

    doc.add_paragraph()

    if down:
        _add_clue_section(doc, "По вертикали ↓", down)


def _add_clue_section(doc: Document, title: str, words: list) -> None:
    """Добавляет секцию вопросов."""
    h = doc.add_heading(title, level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)
        run.font.size = Pt(13)

    for w in words:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)

        # Номер
        run_num = p.add_run(f"{w.number}. ")
        run_num.font.bold = True
        run_num.font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)
        run_num.font.size = Pt(11)

        # Подсказка
        run_hint = p.add_run(w.question.hint)
        run_hint.font.size = Pt(11)

        # Картинка (если есть)
        if w.question.image_path and os.path.exists(w.question.image_path):
            try:
                run_img = p.add_run()
                run_img.add_picture(w.question.image_path, width=Mm(15))
            except Exception:
                pass

        # Категория
        if w.question.category:
            run_cat = p.add_run(f"  [{w.question.category}]")
            run_cat.font.size = Pt(7)
            run_cat.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _set_cell_border(cell, color: str) -> None:
    """Устанавливает рамку ячейки."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def _set_cell_shading(cell, color: str) -> None:
    """Устанавливает фон ячейки."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_margins(cell, top=0, bottom=0, left=0, right=0) -> None:
    """Устанавливает внутренние отступы ячейки (в twips)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(margins)
